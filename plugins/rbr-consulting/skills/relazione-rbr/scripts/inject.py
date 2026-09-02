#!/usr/bin/env python3
"""
RBR Report — inietta header (logo RBR) e footer (testo relazione + numero pagina)
in un .docx generato da helpers.js.

Uso:
    python3 inject.py <input.docx> <output.docx> "Tipo Relazione – Ristorante – Referente"

Dipendenza: python-docx.
"""
import sys
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xE3, 0x06, 0x13)
GREY = RGBColor(0x88, 0x88, 0x88)
ASSETS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "assets")
LOGO = os.path.join(ASSETS, "logo.png")


def _set_border(paragraph, edge="bottom", color="E30613", size="12", space="1"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), size)
    el.set(qn("w:space"), space)
    el.set(qn("w:color"), color)
    pbdr.append(el)


def _add_page_field(paragraph):
    run = paragraph.add_run()
    fldSimple = OxmlElement("w:fldSimple")
    fldSimple.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fldSimple.append(r)
    paragraph._p.append(fldSimple)


def inject(input_path, output_path, footer_text):
    doc = Document(input_path)
    section = doc.sections[0]
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(0.8)

    # -------- HEADER: logo RBR a sinistra + filetto rosso --------
    header = section.header
    header.is_linked_to_previous = False
    # riusa/crea il primo paragrafo
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO):
        run = hp.add_run()
        run.add_picture(LOGO, height=Cm(0.95))
    _set_border(hp, edge="bottom", color="E30613", size="14")

    # -------- FOOTER: testo relazione a sinistra + "pag. N" a destra --------
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ""
    _set_border(fp, edge="top", color="DDDDDD", size="6")
    # tab stop a destra per allineare il numero pagina
    tab_stops = fp.paragraph_format.tab_stops
    from docx.enum.text import WD_TAB_ALIGNMENT
    usable = section.page_width - section.left_margin - section.right_margin
    tab_stops.add_tab_stop(usable, WD_TAB_ALIGNMENT.RIGHT)

    r = fp.add_run(footer_text)
    r.font.size = Pt(7.5)
    r.font.color.rgb = GREY
    r.font.name = "Calibri"

    r2 = fp.add_run("\tpag. ")
    r2.font.size = Pt(7.5)
    r2.font.color.rgb = GREY
    r2.font.name = "Calibri"
    _add_page_field(fp)

    doc.save(output_path)
    print(f"OK -> {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Uso: python3 inject.py <input.docx> <output.docx> \"Footer text\"")
        sys.exit(1)
    inject(sys.argv[1], sys.argv[2], sys.argv[3])
